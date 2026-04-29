import urllib.request
import urllib.error
import json
import time
import re
from docx import Document
from docx.shared import Pt

# ---------- Конфигурация LM Studio API ----------
LM_STUDIO_URL = "http://localhost:1234/v1/chat/completions"

def ask_lm_studio(prompt, system_msg=None, temperature=0.7, max_tokens=12000):
    """Отправляет запрос к LM Studio и возвращает текстовый ответ."""
    headers = {"Content-Type": "application/json"}
    messages = []
    if system_msg:
        messages.append({"role": "system", "content": system_msg})
    messages.append({"role": "user", "content": prompt})

    payload = {
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False
    }
    data = json.dumps(payload).encode('utf-8')

    try:
        req = urllib.request.Request(LM_STUDIO_URL, data=data, headers=headers)
        with urllib.request.urlopen(req, timeout=780) as response:
            response_data = response.read().decode('utf-8')
            result = json.loads(response_data)
            return result["choices"][0]["message"]["content"]
    except Exception as e:
        raise RuntimeError(f"Ошибка при обращении к LM Studio: {e}")

# ---------- Извлечение JSON ----------
def extract_json(text: str):
    """Пытается извлечь JSON-объект из ответа модели."""
    text = re.sub(r'```(?:json)?\s*', '', text, flags=re.IGNORECASE).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1:
            try:
                return json.loads(text[start:end+1])
            except:
                pass
    return None

# ---------- Форматирование Markdown в документ docx ----------
def add_formatted_markdown(doc: Document, text: str):
    """
    Преобразует Markdown-подобный текст в элементы документа docx.
    Поддерживает:
      - Заголовки #, ##, ###
      - **жирный шрифт**
      - Таблицы (строки начинающиеся с | )
      - Обычные абзацы
    """
    lines = text.split('\n')
    buffer = []          # накапливает строки обычного текста
    table_lines = []     # строки текущей таблицы
    in_table = False

    def flush_buffer():
        """Выгружает накопленный текст как отдельные абзацы (по непустым строкам)."""
        for line in buffer:
            line = line.strip()
            if line:
                para = doc.add_paragraph()
                process_inline_formatting(para, line)
        buffer.clear()

    def flush_table():
        """Парсит накопленные строки таблицы и добавляет таблицу в документ."""
        if not table_lines:
            return
        # Очищаем от пустых строк в начале/конце (на всякий случай)
        while table_lines and not table_lines[0].strip():
            table_lines.pop(0)
        while table_lines and not table_lines[-1].strip():
            table_lines.pop()
        if len(table_lines) < 2:
            # слишком короткая таблица – выведем как текст
            for line in table_lines:
                para = doc.add_paragraph(line)
            table_lines.clear()
            return

        # Определяем количество столбцов по первой строке
        header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
        num_cols = len(header_cells)
        # Создаём таблицу: строк = все кроме разделительной (вторая строка)
        data_rows = [line for i, line in enumerate(table_lines) if i != 1]  # 0-заголовок, 1-разделитель, 2+ данные
        table = doc.add_table(rows=len(data_rows), cols=num_cols)
        table.style = 'Table Grid'   # можно заменить на другой стиль, например 'Light Shading Accent 1'
        # Заполняем заголовки
        for j, cell_text in enumerate(header_cells):
            cell = table.cell(0, j)
            cell.text = ''
            p = cell.paragraphs[0]
            process_inline_formatting(p, cell_text)
        # Заполняем данные
        for i, row_line in enumerate(data_rows[1:], start=1):
            row_cells = [c.strip() for c in row_line.split('|')[1:-1]]
            for j, cell_text in enumerate(row_cells):
                if j < num_cols:
                    cell = table.cell(i, j)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    process_inline_formatting(p, cell_text)
        # Добавляем пустой абзац после таблицы для отступа
        doc.add_paragraph('')
        table_lines.clear()

    # Главный цикл разбора строк
    for line in lines:
        stripped = line.strip()

        # Переключение в режим таблицы (строка начинается с |)
        if stripped.startswith('|'):
            if in_table:
                table_lines.append(line)
            else:
                flush_buffer()          # сбросить накопленный текст
                in_table = True
                table_lines.append(line)
            continue

        # Если мы были в таблице, а теперь строка не табличная – завершаем таблицу
        if in_table:
            flush_table()
            in_table = False
            # текущую строку будем обрабатывать как обычную (ниже)

        # Пустая строка – разделитель абзацев
        if not stripped:
            flush_buffer()
            continue

        # Заголовки Markdown
        if re.match(r'^#{1,3}\s', stripped):
            flush_buffer()
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = stripped[level:].strip()
            doc.add_heading(heading_text, level=level)
            continue

        # Обычная текстовая строка
        buffer.append(line)

    # Обработка остатков
    if in_table:
        flush_table()
    else:
        flush_buffer()

def process_inline_formatting(paragraph, line: str):
    """Добавляет runs с учётом **жирного** текста."""
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last_end = 0
    for match in pattern.finditer(line):
        start, end = match.span()
        # Текст до жирного
        if start > last_end:
            normal_text = line[last_end:start]
            run = paragraph.add_run(normal_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        # Жирный фрагмент
        bold_text = match.group(1)
        run = paragraph.add_run(bold_text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        last_end = end
    # Остаток строки
    if last_end < len(line):
        normal_text = line[last_end:]
        run = paragraph.add_run(normal_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

# ---------- Ввод от пользователя ----------
topic = input("Введите тему дипломной работы: ").strip()
edits = input("Введите дополнительные правки / требования: ").strip()

# Проверка связи
try:
    test = ask_lm_studio("ping", system_msg="Ответь одним словом: pong")
    print("Связь с LM Studio установлена.\n")
except Exception as e:
    print(f"Ошибка подключения: {e}")
    exit(1)

# ---------- Шаг 1: Получение подробного плана ----------
plan_prompt = f"""
Ты — академический консультант. Составь детальный план дипломной работы на тему:
"{topic}"

Дополнительные требования: {edits}

План должен включать главы и их подпункты (разделы, подразделы). Выдай результат СТРОГО в формате JSON без лишнего текста.
Формат JSON:
{{
  "Глава 1. Название главы": {{
    "description": "Краткое описание главы (1 предложение)",
    "subsections": {{
      "1.1 Название подпункта": "Что раскрыть в этом подпункте",
      "1.2 Название подпункта": "...",
      "1.2.1 Подподпункт": "..."
    }}
  }},
  ...
}}
Не добавляй никаких комментариев, только JSON.
JSON:
"""

print("Запрашиваю детальный план...")
plan_response = ask_lm_studio(plan_prompt, system_msg="Ты — полезный ассистент, который всегда отвечает строго по инструкции.")
plan = extract_json(plan_response)

if not plan:
    print("Не удалось распарсить план. Ответ модели:")
    print(plan_response)
    exit(1)

print("\nПолученный план:")
for chapter, data in plan.items():
    print(f"\n{chapter}: {data['description']}")
    for sub, desc in data.get('subsections', {}).items():
        print(f"   {sub}: {desc}")

# ---------- Шаг 2: Генерация глав с подпунктами ----------
document = Document()
document.styles['Normal'].font.name = 'Times New Roman'
document.styles['Normal'].font.size = Pt(14)

document.add_heading('Дипломная работа', level=0)
document.add_paragraph(f'Тема: {topic}')
if edits:
    document.add_paragraph(f'Дополнительные правки: {edits}')
document.add_paragraph('')

for chapter, data in plan.items():
    print(f"\nГенерирую: {chapter} ...")
    description = data['description']
    subsections = data.get('subsections', {})

    # Собираем описание всех подпунктов для промпта
    subs_text = ""
    for key, desc in subsections.items():
        subs_text += f"- {key}: {desc}\n"

    chapter_prompt = f"""
Ты — опытный автор дипломных работ. Напиши текст для главы "{chapter}" дипломной работы.

Общая тема: "{topic}"
Дополнительные требования: {edits}

Описание главы: {description}

В этой главе должны быть раскрыты следующие подпункты:
{subs_text}

Напиши полный, связный текст главы. Обязательно используй заголовки подпунктов в формате:
## 1.1 Название подпункта
или
### 1.2.1 Название подподпункта
(в зависимости от уровня).

Текст должен быть академическим, подробным, с логическими переходами. Начинай сразу с первого подпункта, не дублируя общий заголовок главы.
Текст главы:
"""
    try:
        chapter_text = ask_lm_studio(
            chapter_prompt,
            system_msg="Ты — академический писатель.",
            max_tokens=8000  # достаточно для главы со всеми подпунктами
        )
    except Exception as e:
        print(f"Ошибка при генерации '{chapter}': {e}")
        chapter_text = f"[Ошибка: {e}]"

    # Добавляем общий заголовок главы в документ
    document.add_heading(chapter, level=1)

    # Форматируем Markdown от модели (уже содержит ## и т.д.)
    add_formatted_markdown(document, chapter_text)

    time.sleep(1.0)  # небольшая пауза между запросами

# ---------- Сохранение ----------
output_file = "Дипломная_работа.docx"
document.save(output_file)
print(f"\n✅ Документ сохранён: {output_file}")